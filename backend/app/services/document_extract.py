"""从上传的简历文档中提取纯文本。

支持格式：
- PDF  (.pdf)  -> pypdf
- Word (.docx) -> python-docx
- 纯文本 (.txt/.md/.text) -> 直接解码（utf-8，回退 gbk/latin-1）

设计要点：
- 纯同步、无外部 IO，便于在 FastAPI 同步端点中直接调用（端点由 FastAPI 线程池调度）。
- 提取失败时抛出 ExtractionError，由路由层转成友好 HTTP 错误。
- 图片型/扫描件 PDF（无文本层）会提取为空，这里会明确提示，而不是静默返回空串。
"""
from __future__ import annotations

import io
import os

from pydantic import BaseModel


class ExtractionError(Exception):
    """文档解析失败（不支持的格式 / 无文本层 / 损坏文件等）。"""


class ExtractionResult(BaseModel):
    text: str
    pages: int
    fileName: str
    charCount: int
    warning: str | None = None


_ALLOWED_EXT = {".pdf", ".docx", ".doc", ".txt", ".md", ".text"}


def _read_pdf(content: bytes) -> ExtractionResult:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages_text: list[str] = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages_text.append("")
    text = "\n\n".join(t for t in pages_text if t.strip())
    warning = None
    if not text.strip():
        warning = "未能从 PDF 中提取到文本，可能是扫描件（图片型）或加密文件，请改用可复制文字的 PDF 或手动粘贴文本。"
    return ExtractionResult(
        text=text,
        pages=len(reader.pages),
        fileName="",
        charCount=len(text),
        warning=warning,
    )


def _read_docx(content: bytes) -> ExtractionResult:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # 一并抽取表格文本，简历常见
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    warning = None
    if not text.strip():
        warning = "未能从 Word 文档中提取到文本，请确认文件未损坏。"
    return ExtractionResult(
        text=text,
        pages=len(doc.paragraphs),
        fileName="",
        charCount=len(text),
        warning=warning,
    )


def _read_plain(content: bytes) -> ExtractionResult:
    text = None
    for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            text = content.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if text is None:
        text = content.decode("utf-8", errors="replace")
    lines = [ln for ln in text.splitlines() if ln.strip()]
    body = "\n".join(lines)
    return ExtractionResult(
        text=body,
        pages=0,
        fileName="",
        charCount=len(body),
        warning=None,
    )


def extract_text(filename: str, content: bytes) -> ExtractionResult:
    """根据扩展名分流提取文本。"""
    if not content:
        raise ExtractionError("文件内容为空。")

    ext = os.path.splitext(filename or "")[1].lower()
    if ext not in _ALLOWED_EXT:
        raise ExtractionError(
            f"不支持的文件格式：{ext or '未知'}。请上传 PDF、Word(.docx) 或纯文本(.txt/.md) 简历。"
        )

    try:
        if ext == ".pdf":
            result = _read_pdf(content)
        elif ext in (".docx", ".doc"):
            result = _read_docx(content)
        else:
            result = _read_plain(content)
    except ExtractionError:
        raise
    except Exception as e:  # noqa: BLE001
        raise ExtractionError(f"解析文件失败：{e}") from e

    result.fileName = filename
    return result
